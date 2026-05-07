"""
Canonical styling helper for BrainWise closeout documents.

write_doc(path, title, subtitle, sections) generates a binary .docx file
from a list of (style, text) tuples.

Supported styles:
  h1       - top-level section heading (16pt bold)
  h2       - subsection heading (13pt bold)
  h3       - sub-subsection heading (11pt bold)
  body     - paragraph text
  bullet   - bulleted list item
  code     - monospace block (Consolas 10pt)
  spacer   - empty paragraph
"""
from docx import Document
from docx.shared import Pt, RGBColor


def write_doc(path, title, subtitle, sections):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    tr = t.add_run(title)
    tr.bold = True
    tr.font.size = Pt(20)

    s = doc.add_paragraph()
    sr = s.add_run(subtitle)
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for style_name, text in sections:
        if style_name == 'h1':
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(16)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif style_name == 'h2':
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif style_name == 'h3':
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
        elif style_name == 'body':
            doc.add_paragraph(text)
        elif style_name == 'bullet':
            doc.add_paragraph(text, style='List Bullet')
        elif style_name == 'code':
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        elif style_name == 'spacer':
            doc.add_paragraph()

    doc.save(path)
    return path
